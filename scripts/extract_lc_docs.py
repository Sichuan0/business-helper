#!/usr/bin/env python3
"""
从一份含多张单证表格的 docx（如教师答案 / 练习册）中，按表格索引拆出独立 docx，
并应用文本修正（HS 编码、笔误等）。

用法:
    python extract_lc_docs.py --src "答案.docx" --out-dir ./输出
    python extract_lc_docs.py --src "答案.docx" --out-dir ./输出 --keep 1,2,3,5,7

默认拆出: 1=商业发票 2=装箱单 3=Form A 原产地证 4=报检单 5=保险单 7=汇票
  (索引对应原 docx 里表格的顺序，0 通常是题目资料块、6 通常是提单)

修正规则见 FIXES，按你自己的 L/C 改。
"""
import argparse
import sys
from pathlib import Path
from docx import Document

# 默认拆出的表格索引 -> 输出文件名前缀
DEFAULT_TARGETS = {
    1: "01_商业发票_COMMERCIAL_INVOICE",
    2: "02_装箱单_PACKING_LIST",
    3: "03_原产地证_GSP_FORM_A",
    4: "04_报检单",
    5: "05_保险单_INSURANCE_POLICY",
    7: "07_汇票_BILL_OF_EXCHANGE",
}

# 文本修正（find -> replace），按需修改
FIXES = [
    ("560KCS", "560KGS"),          # 装箱单毛重笔误示例
    ("5911320000", "6301.0000"),   # HS 编码示例
]


def apply_fixes(table):
    fixed = 0
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                # run 级替换，保留格式
                for run in para.runs:
                    for find, repl in FIXES:
                        if find in run.text:
                            run.text = run.text.replace(find, repl)
                            fixed += 1
                # 兜底：文本跨 run 时按整段重建
                full = para.text
                if any(f in full for f, _ in FIXES):
                    new = full
                    for find, repl in FIXES:
                        new = new.replace(find, repl)
                    if para.runs:
                        para.runs[0].text = new
                        for r in para.runs[1:]:
                            r.text = ""
                    fixed += 1
    return fixed


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

    ap = argparse.ArgumentParser(description="按表格索引拆分 docx 单证并做文本修正")
    ap.add_argument("--src", required=True, help="输入 docx 路径")
    ap.add_argument("--out-dir", required=True, help="输出目录")
    ap.add_argument("--keep", default="", help="要保留的表格索引，逗号分隔，如 1,2,3,5,7；省略=用默认")
    args = ap.parse_args()

    if args.keep:
        targets = {int(i): DEFAULT_TARGETS.get(int(i), f"{int(i):02d}_table{int(i)}")
                   for i in args.keep.split(",")}
    else:
        targets = dict(DEFAULT_TARGETS)

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    for idx, name in targets.items():
        doc = Document(args.src)
        tables = list(doc.tables)
        keep = None
        for i, t in enumerate(tables):
            if i == idx:
                keep = t
            else:
                t._element.getparent().remove(t._element)
        if keep is None:
            print(f"[warn] 表 {idx} 未找到（该 docx 只有 {len(tables)} 张表）")
            continue
        # 删除所有段落（题目原文、装船通知等），只留目标表格
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
        n = apply_fixes(keep)
        out = out_dir / f"{name}.docx"
        doc.save(out)
        print(f"saved: {out.name}  (修正 {n} 处)")
    print("done")


if __name__ == "__main__":
    main()
