#!/usr/bin/env python3
"""
trade-doc 渲染脚本：把 filled JSON 渲染成外贸单证 docx。

用法:
    python render_doc.py --data filled.json --out PI-2026-001.docx

filled JSON 结构: flat, 顶层带 doc_type。字段定义见 schemas/<doc_type>.json。
脚本只负责确定性渲染；抽取与追问由 Claude 在对话中完成。
"""
import argparse
import json
import sys
from datetime import date
from pathlib import Path

# Windows 控制台默认 GBK，强制 stdout/stderr 用 UTF-8，避免中文乱码影响读取
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------- 通用工具 ----------

def money(amount, currency=""):
    """格式化金额: 1,234.50 带千分位; 可选币种前缀。"""
    try:
        s = f"{float(amount):,.2f}"
    except (TypeError, ValueError):
        s = str(amount) if amount is not None else ""
    return f"{currency} {s}".strip() if currency else (s or "")


def set_cell(cell, text, bold=False, align=None, size=None):
    """写单元格内容并设置样式。"""
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    if size:
        run.font.size = Pt(size)


def add_kv_line(doc, label, value):
    """加一行「标签: 值」，值为空则显示 —。"""
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(value if value else "—")


# ---------- PI 渲染 ----------

# 必填字段（pi_date / total_amount / items[].total_price 为自动计算，不在此列）
PI_REQUIRED = [
    "pi_number", "seller_name", "buyer_name",
    "incoterms", "currency", "payment_terms", "items",
]


def render_proforma_invoice(data, out_path):
    # 1. 校验必填（一次性收集顶层 + 明细的全部缺失）
    items = data.get("items") or []
    item_req = ["description", "quantity", "unit", "unit_price"]
    missing = [f for f in PI_REQUIRED if not data.get(f)]
    item_missing = []
    for i, it in enumerate(items, 1):
        miss = [f for f in item_req if it.get(f) in (None, "", [])]
        if miss:
            item_missing.append(f"第{i}项缺{miss}")
    if missing or item_missing:
        sys.exit(f"[render] 缺必填字段: {missing + item_missing}。请补齐后重试。")

    currency = data.get("currency", "")

    # 2. 自动计算
    for it in items:
        if it.get("quantity") is not None and it.get("unit_price") is not None:
            try:
                it["total_price"] = round(float(it["quantity"]) * float(it["unit_price"]), 2)
            except (TypeError, ValueError):
                pass
    try:
        data["total_amount"] = round(sum(float(it.get("total_price") or 0) for it in items), 2)
    except (TypeError, ValueError):
        data["total_amount"] = data.get("total_amount", "")

    pi_date = data.get("pi_date") or str(date.today())

    # 3. 构建文档
    doc = Document()

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PROFORMA INVOICE")
    r.bold = True
    r.font.size = Pt(18)

    # PI 号 + 日期
    hdr = doc.add_paragraph()
    hdr.add_run("PI No.: ").bold = True
    hdr.add_run(data.get("pi_number", ""))
    hdr.add_run("        ")
    hdr.add_run("Date: ").bold = True
    hdr.add_run(pi_date)

    # 卖方 / 买方（无边框 2 列表）
    party = doc.add_table(rows=0, cols=2)
    party_left = ["SELLER", data.get("seller_name", ""), data.get("seller_address", ""),
                  data.get("seller_contact", "")]
    party_right = ["BUYER", data.get("buyer_name", ""), data.get("buyer_address", ""),
                   data.get("buyer_contact", "")]
    if data.get("seller_tax_id"):
        party_left.append(f"Tax ID: {data['seller_tax_id']}")
        party_right.append(f"Tax ID: {data['buyer_tax_id']}" if data.get("buyer_tax_id") else "")
    elif data.get("buyer_tax_id"):
        party_left.append("")
        party_right.append(f"Tax ID: {data['buyer_tax_id']}")
    for i in range(len(party_left)):
        cells = party.add_row().cells
        set_cell(cells[0], party_left[i], bold=(i == 0))
        set_cell(cells[1], party_right[i], bold=(i == 0))

    doc.add_paragraph()

    # 产品明细表
    doc.add_paragraph().add_run("Product Details").bold = True
    cols = ["#", "Description", "HS Code", "Qty", "Unit", "Unit Price", "Total Price"]
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cols):
        set_cell(tbl.rows[0].cells[i], c, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

    for idx, it in enumerate(items, 1):
        cells = tbl.add_row().cells
        set_cell(cells[0], idx, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
        set_cell(cells[1], it.get("description", ""), size=9)
        set_cell(cells[2], it.get("hs_code", ""), align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
        set_cell(cells[3], it.get("quantity", ""), align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
        set_cell(cells[4], it.get("unit", ""), align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
        set_cell(cells[5], money(it.get("unit_price"), currency),
                 align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
        set_cell(cells[6], money(it.get("total_price"), currency),
                 align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)

    # 合计行（合并前 6 列为 "TOTAL"）
    total_row = tbl.add_row()
    merged = total_row.cells[0].merge(total_row.cells[5])
    set_cell(merged, "TOTAL", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
    set_cell(total_row.cells[6], money(data.get("total_amount"), currency),
             bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)

    doc.add_paragraph()

    # 条款
    doc.add_paragraph().add_run("Terms & Conditions").bold = True
    add_kv_line(doc, "Trade Terms (Incoterms)", data.get("incoterms"))
    add_kv_line(doc, "Payment Terms", data.get("payment_terms"))
    add_kv_line(doc, "Delivery Time", data.get("delivery_time"))
    add_kv_line(doc, "Validity", data.get("validity"))
    add_kv_line(doc, "Country of Origin", data.get("origin"))
    add_kv_line(doc, "Shipping Method", data.get("shipping_method"))

    if data.get("remarks"):
        p = doc.add_paragraph()
        p.add_run("Remarks: ").bold = True
        p.add_run(data.get("remarks"))

    # 签字栏
    doc.add_paragraph()
    sig = doc.add_table(rows=2, cols=2)
    set_cell(sig.rows[0].cells[0], "_______________________")
    set_cell(sig.rows[0].cells[1], "_______________________")
    set_cell(sig.rows[1].cells[0], "Seller", bold=True)
    set_cell(sig.rows[1].cells[1], "Buyer", bold=True)

    # 4. 保存
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"[render] 已生成: {out.resolve()}")
    print(f"[render] 合计 {money(data.get('total_amount'), currency)} | {len(items)} 项明细")


# ---------- 入口 ----------

RENDERERS = {
    "proforma_invoice": render_proforma_invoice,
}


def main():
    ap = argparse.ArgumentParser(description="trade-doc 渲染: filled JSON -> docx")
    ap.add_argument("--data", required=True, help="filled JSON 路径")
    ap.add_argument("--out", help="输出 docx 路径; 省略则用 pi_number 命名到当前目录")
    args = ap.parse_args()

    data = json.loads(Path(args.data).read_text(encoding="utf-8"))
    doc_type = data.get("doc_type", "proforma_invoice")
    renderer = RENDERERS.get(doc_type)
    if not renderer:
        sys.exit(f"[render] 不支持的单证类型: {doc_type}。支持: {list(RENDERERS)}")

    out_path = args.out
    if not out_path:
        name = data.get("pi_number") or doc_type
        out_path = f"{name}.docx"
    renderer(data, out_path)


if __name__ == "__main__":
    main()
